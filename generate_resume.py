import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text):
    heading = doc.add_heading(text, level=1)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121) # Sleek deep blue
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(6)

def add_job(doc, title, company, location, date, descriptions):
    # Company and Location
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run_comp = p.add_run(f"{company} ")
    run_comp.bold = True
    run_comp.font.size = Pt(11)
    
    if location:
        run_loc = p.add_run(f"| {location}")
        run_loc.font.size = Pt(11)
    
    # Title and Date
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    run_title = p2.add_run(title)
    run_title.italic = True
    run_title.font.size = Pt(11)
    
    run_date = p2.add_run(f" ({date})")
    run_date.font.size = Pt(11)

    for desc in descriptions:
        p_desc = doc.add_paragraph(style='List Bullet')
        p_desc.paragraph_format.left_indent = Inches(0.25)
        p_desc.paragraph_format.space_after = Pt(4)
        run_d = p_desc.add_run(desc)
        run_d.font.size = Pt(10.5)

def main():
    doc = docx.Document()

    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Change Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    # Header / Contact Info
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(4)
    name_run = header_p.add_run("JOHNATHAN D. YORK, PMP, PMI-CPMAI\n")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(31, 78, 121)

    contact_run = header_p.add_run("Woodstock, MD 21163 | (843) 754-7350 | Jyork143@gmail.com\nLinkedIn: linkedin.com/in/johnathan-york")
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(100, 100, 100)

    # Professional Profile
    add_heading(doc, "PROFESSIONAL PROFILE")
    p_profile = doc.add_paragraph("Highly accomplished Senior Project Manager with over 15 years of elite experience driving large-scale capital projects, engineering operations, and infrastructure modernization across heavily regulated nuclear and energy architectures. A U.S. Navy Nuclear Program veteran with an exceptional mastery of nuclear plant operations, safety culture, and rigid quality governance frameworks. Certified PMP with a proven track record managing multi-million dollar budgets, cross-functional engineering teams, and strict regulatory compliance (including JFMM and nuclear PPEP). Adept at pioneering advanced technology integrations and cognitive workflow automations to safely optimize scheduling, reduce down-time, and enforce zero-tolerance risk management in mission-critical environments.")
    p_profile.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Core Competencies
    add_heading(doc, "CORE COMPETENCIES")
    competencies = [
        "Capital Project Planning & Execution Process (PPEP)",
        "Regulated Nuclear Environment Governance & Safety Culture",
        "Multi-Million Dollar Infrastructure & Capital Budget Management",
        "Quality Assurance Frameworks (Joint Fleet Maintenance Manual - JFMM)",
        "Root Cause Analysis, Risk Mitigation & Failure Mode Assessments",
        "Resource Allocation, Earned Value Management & Vendor Oversight",
        "Advanced Technology Integration & Workflow Digital Transformation",
        "Cross-Functional Leadership & High-Stakes Stakeholder Management"
    ]
    
    # Create a 2-column table for competencies
    table = doc.add_table(rows=4, cols=2)
    table.autofit = True
    for i in range(4):
        p1 = table.cell(i, 0).paragraphs[0]
        p1.style = 'List Bullet'
        p1.paragraph_format.left_indent = Inches(0.25)
        p1.add_run(competencies[i*2]).font.size = Pt(10)
        
        p2 = table.cell(i, 1).paragraphs[0]
        p2.style = 'List Bullet'
        p2.paragraph_format.left_indent = Inches(0.25)
        p2.add_run(competencies[i*2+1]).font.size = Pt(10)
    
    # Professional Experience
    add_heading(doc, "PROFESSIONAL EXPERIENCE")

    add_job(doc, "Nuclear Project Manager, Capital Projects", "XCEL ENERGY (via SelectSource Int'l)", "Prairie Island Nuclear Plant, MN", "2025 - 2025", [
        "Led capital project planning and project execution support under the strict, regulated parameters of a commercial nuclear power facility.",
        "Managed complex project workflows, risk matrices, and data governance alignments supporting the site's Project Planning and Execution Process (PPEP).",
        "Evaluated heavy engineering and operational workflows to isolate cognitive vs. non-cognitive tasks, mapping optimal entry points for data-driven tools to enhance team productivity and decision support while maintaining absolute safety overhead.",
        "Pioneered AI adoption frameworks and governed training deployment for Microsoft Copilot, enabling legacy technical teams to safely leverage advanced computing within regulatory boundaries and limitations.",
        "Evaluated emerging technologies against rigorous corporate risk, data privacy, and operational readiness metrics utilizing foundational PMI-CPMAI™ frameworks."
    ])

    add_job(doc, "Program Manager, North America Operational Excellence", "AMAZON", "North America", "2021 - 2024", [
        "Directed overarching project installation, critical infrastructure logistics, and network operations for complex automated logistics networks across North America.",
        "Managed international resource allocations, vendor contracts, and risk response strategies for a $50M annual operations budget, delivering programmatic cost savings that outperformed target goals by 42%.",
        "Automated comprehensive asset analysis and predictive health metrics for critical regional electrical power and network connectivity systems, reducing product downtime by 40%.",
        "Utilized cloud-powered digital augmentations and modernized fleet mapping systems to drive a 37% increase in hardware infrastructure activations while cutting installation failure rates by 52.1% YoY.",
        "Revamped workflow automation pipelines to eliminate operational backlogs, lifting on-site delivery and scheduling accuracy by 23% YoY to align with enterprise performance metrics."
    ])

    add_job(doc, "Engineering Project Manager, Intelligent Tools", "STARBUCKS TRYER LABORATORY (via Ascendion)", "Seattle, WA", "2024 - 2025", [
        "Spearheaded the full-lifecycle development, risk management, and strategic deployment of data-centric tools and custom automation applications to optimize enterprise asset and infrastructure tracking.",
        "Deployed a centralized Maintenance Management System and automated Inventory Tracking Tool that increased real-time operational visibility by 80% and reduced YoY corporate audit gaps by 45%.",
        "Developed customized electrical distribution and infrastructure telemetry tools that accelerated time-to-deployment by 30% across cross-functional facilities.",
        "Enforced rigid technology governance and validation frameworks, establishing explicit corporate metrics for risk detection, data privacy compliance, transparency, and data integrity."
    ])

    add_job(doc, "Site Operations Manager, Flagship Test Center", "KIWA PVEL", "South San Francisco, CA", "2019 - 2021", [
        "Recruited to greenfield-onboard, staff, and execute operations for a high-voltage testing facility, managing its strategic scale-up into a multi-megawatt Inverter and Battery Energy Storage Systems (BESS) lab to anchor flagship Product Qualification Programs (PQP).",
        "Commissioned and supervised the complex engineering installation of two separate independent AC grid simulators, facilitating empirical bankability testing under simulated grid anomalies and dynamic utility conditions.",
        "Directed daily operations for extended-reliability testing, including environmental chamber testing, Incidence Angle Modifier (IAM), and High-Throughput Module Inspections (HTMI) under rigid safety standards.",
        "Maintained 100% operational continuity and testing throughput while managing laboratory personnel through the height of the COVID-19 pandemic and a major international corporate acquisition by Kiwa."
    ])

    add_job(doc, "Project Lead & Autonomous Systems Researcher", "NASA STUDENT LAUNCH PROGRAM / UNIVERSITY OF HAWAII", "Honolulu, HI", "2017 - 2019", [
        "Concepted, engineered, and built a fully autonomous, solar-powered smart irrigation system designed to restore native vegetation and mitigate critical soil erosion on the island of Kaho'olawe; presented findings at the national SACNAS Conference (San Antonio, TX).",
        "Formulated and executed real-time Python-based control algorithms and telemetry logic for autonomous platforms, successfully decreasing system latency by 40% and ensuring 100% prototype field reliability.",
        "Led competitive engineering cohorts in the elite ARLISS aerospace competition in Black Rock, Nevada, commanding the design and deployment of autonomous satellites under 100% NASA flight-readiness compliance."
    ])

    add_job(doc, "Nuclear Electrician / Submarine Warfare Specialist", "UNITED STATES NAVY | Submarine Force", "USS Columbia (SSN 771) & MARF Prototype", "2009 - 2016", [
        "Supervised, maintained, and operated mission-critical nuclear propulsion and electrical generation plant infrastructure under strict Department of Energy (DOE) and military operational tolerances.",
        "Selected as a Junior Staff Instructor at the MARF DOE nuclear training prototype based on technical aptitude; qualified as a Shutdown Reactor Operator (SRO), Throttleman, Craftsman, and Dual-Plant Cleanliness Inspector (Reactor/Secondary), training future fleet operators.",
        "Partnered directly with the Engineering Department Master Chief to execute a comprehensive operational revamp of the submarine's engineering department; qualified as Quality Assurance Inspector (QAI), QA Planner (QAP), and QA Work Center Supervisor under the rigid structural governance of the Joint Fleet Maintenance Manual (JFMM).",
        "Served as the primary Ship's Force Representative for complex shipyard overhaul periods due to expertise in auditing and executing highly regulated technical work packages and technical requirements.",
        "Crew recipient of the Arleigh Burke Fleet Trophy, Battle Effectiveness Award (Battle \"E\"), and Engineering Excellence Award (Engineering \"E\") across multiple classified national security deployments and a maximum-grade Operational Reactor Safeguards Examination (ORSE).",
        "Earned the U.S. Navy Submarine Warfare Designation (\"Dolphins\") in an unprecedented 3 months, dramatically outperforming the standard 12-to-18-month qualification timeline."
    ])

    # Education & Credentials
    add_heading(doc, "EDUCATION & CREDENTIALS")
    
    edu = [
        "University of Maryland Global Campus - Bachelor of Science in AI Applications (In Progress)",
        "University of Maryland Global Campus - Associate of Arts in Business and Management",
        "US Naval Nuclear Power School - Graduate (Nuclear Power Plant Construction & Reactor Theory)",
        "US Naval Nuclear Training Prototype - Graduate (Nuclear Plant Operation & Maintenance)"
    ]
    for e in edu:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(e).font.size = Pt(10.5)

    p_cert_title = doc.add_paragraph("\nCertifications & Professional Credentials:")
    p_cert_title.runs[0].bold = True
    
    certs = [
        "Project Management Professional (PMP)® - Project Management Institute (Credential ID: 3959827)",
        "PMI Certified Professional in Managing AI (PMI-CPMAI)™ - Project Management Institute",
        "PMI Advanced Micro-Credentials:"
    ]
    for c in certs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(c).font.size = Pt(10.5)

    sub_certs = [
        "Practical Application of GenAI for Project Managers",
        "AI in Infrastructure & Construction Management",
        "Talking to AI (Prompt Engineering)",
        "Introduction to Cognitive Project Management in AI"
    ]
    for sc in sub_certs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(f"- {sc}").font.size = Pt(10.5)

    doc.save("Johnathan_York_Resume.docx")

if __name__ == "__main__":
    main()
